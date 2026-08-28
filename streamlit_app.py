
import io
import re
from pathlib import Path

import numpy as np
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
import matplotlib.pyplot as plt


st.set_page_config(
    page_title="OBE Evaluation Report Generator",
    page_icon="📊",
    layout="wide"
)

# -----------------------------
# Styling
# -----------------------------
st.markdown("""
<style>
.main-title {
    font-size: 42px;
    font-weight: 800;
    color: #17365D;
    margin-bottom: 2px;
}
.subtitle {
    color: #6B7280;
    font-size: 17px;
    margin-bottom: 28px;
}
.section-title {
    color: #17365D;
    font-size: 30px;
    font-weight: 750;
    margin-top: 28px;
    margin-bottom: 10px;
}
.info-card {
    background: #F4F7FB;
    border: 1px solid #D9E2F3;
    border-radius: 12px;
    padding: 14px 18px;
    min-height: 88px;
    margin-bottom: 8px;
}
.info-label {
    color: #64748B;
    font-size: 14px;
}
.info-value {
    color: #172033;
    font-size: 17px;
    font-weight: 600;
}
.small-note {
    color: #64748B;
    font-size: 13px;
}
</style>
""", unsafe_allow_html=True)

st.markdown(
    '<div class="main-title">📊 OBE Evaluation Report Generator</div>',
    unsafe_allow_html=True
)
st.markdown(
    '<div class="subtitle">Generate a complete, auditable CLO analysis with Word report, Excel analysis workbook and three separate bar-chart PNG files.</div>',
    unsafe_allow_html=True
)

# -----------------------------
# Constants
# -----------------------------
MISSING = "Not available in the provided files."
NO_EVIDENCE = "No assessment evidence for this CLO was identified in the provided Excel file."

with st.sidebar:
    st.header("Report Settings")
    benchmark = st.number_input(
        "OBE benchmark (%)",
        min_value=0.0,
        max_value=100.0,
        value=70.0,
        step=1.0
    )
    st.caption("Used for CLO benchmark achievement and student benchmark calculations.")
    st.divider()
    st.subheader("Input")
    st.caption("Upload the Course Outline and the original OBE assessment Excel workbook. Course information and CLOs can also be entered manually.")

BENCHMARK = float(benchmark)


def clean(v):
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    return str(v).replace("\ufeff", "").replace("\xa0", " ").strip()


def safe(v):
    return clean(v) if clean(v) else MISSING


def num(v):
    try:
        return float(v)
    except Exception:
        return np.nan


def normalize_clo(v):
    m = re.search(r"\bCLO\s*[-_]?\s*(\d+)\b", clean(v), re.I)
    return f"CLO{int(m.group(1))}" if m else ""


def status(x):
    if pd.isna(x):
        return MISSING
    if x >= 80:
        return "Strong"
    if x >= 70:
        return "Satisfactory"
    return "Needs Improvement"


def parse_outline(data):
    """Extract course information, objectives and exact CLO wording from a DOCX outline."""
    doc = Document(io.BytesIO(data))

    paragraph_lines = [clean(p.text) for p in doc.paragraphs if clean(p.text)]

    table_cells = []
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            vals = [clean(c.text) for c in row.cells]
            table_rows.append(vals)
            for v in vals:
                if v:
                    table_cells.append(v)

    # Keep both paragraph and table text because many university course outlines
    # store metadata/CLOs inside tables.
    all_lines = paragraph_lines + table_cells
    text = "\n".join(all_lines)

    info = {
        "Institution": "",
        "Department": "",
        "Program": "",
        "Course Title": "",
        "Course Code": "",
        "Semester": "",
        "Academic Year": "",
        "Campus": "",
        "Instructor/Faculty": "",
        "Credit Hours": "",
        "Section": "",
        "Course Description": ""
    }

    # Label -> adjacent value in a DOCX table.
    label_map = {
        "Course": "Course Title",
        "Course code": "Course Code",
        "Year/Semester": "Semester",
        "Program": "Program",
        "Units/Cr Hrs.": "Credit Hours",
        "Units/Cr Hrs": "Credit Hours",
        "Department": "Department",
        "Instructor": "Instructor/Faculty",
        "Teacher": "Instructor/Faculty",
        "Faculty": "Instructor/Faculty",
        "Academic Year": "Academic Year",
        "Section": "Section",
        "Campus": "Campus"
    }

    for row in table_rows:
        for i, cell in enumerate(row):
            label = clean(cell).rstrip(":")
            if label in label_map and i + 1 < len(row):
                value = clean(row[i + 1])
                if value:
                    info[label_map[label]] = value

    # Conventional text-line patterns as a fallback.
    patterns = {
        "Course Title": r"Course:\s*([^\n]+)",
        "Course Code": r"Course\s*code:\s*([^\s\n]+)",
        "Semester": r"Year/Semester:\s*([^\n]+)",
        "Program": r"Program:\s*([^\n]+)",
        "Credit Hours": r"Units/Cr\s*Hrs\.?:\s*([^\n]+)",
        "Instructor/Faculty": r"(?:Instructor|Teacher|Faculty)\s*:?\s*([^\n]*)",
        "Department": r"Department:\s*([^\n]+)",
        "Academic Year": r"Academic Year:\s*([^\n]+)",
        "Section": r"Section:\s*([^\n]+)"
    }
    for key, pattern in patterns.items():
        m = re.search(pattern, text, re.I)
        if m and not info[key]:
            info[key] = clean(m.group(1))

    # Institution/campus commonly appear as standalone paragraphs.
    if all_lines:
        has_fast = any("FAST School of Management" in x for x in all_lines)
        has_nuces = any("National University of Computer" in x for x in all_lines)
        if has_fast and has_nuces:
            info["Institution"] = "FAST School of Management, National University of Computer & Emerging Sciences"
        elif has_fast:
            info["Institution"] = "FAST School of Management"
        elif has_nuces:
            info["Institution"] = "National University of Computer & Emerging Sciences"
        for x in all_lines:
            mcamp = re.search(r"\b(Lahore Campus|Islamabad Campus|Karachi Campus|Peshawar Campus)\b", x, re.I)
            if mcamp:
                info["Campus"] = clean(mcamp.group(1))
                break

    if info["Instructor/Faculty"] and info["Instructor/Faculty"].rstrip(":").lower() in {"course type", "support", "email", "phone"}:
        info["Instructor/Faculty"] = ""

    # Course description from paragraphs.
    m = re.search(
        r"COURSE DESCRIPTION\s*(.*?)(?:Program Educational Objectives|Course Objectives|Program Learning Outcome|Course Learning Outcomes)",
        text, re.I | re.S
    )
    if m:
        info["Course Description"] = " ".join(m.group(1).split())

    # Objectives: the FAST outline stores them in the table whose first row\n    # begins "Upon successful completion...". The objective text is the third cell.\n    objectives = []\n    for table in doc.tables:\n        rows = table.rows\n        marker = False\n        for row in rows:\n            vals = [clean(c.text) for c in row.cells]\n            if vals and "Upon successful completion of the course" in vals[0]:\n                marker = True\n                continue\n            if marker:\n                if vals and re.fullmatch(r"\\d+", clean(vals[0])):\n                    candidate = clean(vals[2] if len(vals) >= 3 else vals[1] if len(vals) >= 2 else "")\n                    if candidate:\n                        objectives.append(candidate)\n                elif vals:\n                    break\n        if objectives:\n            break\n\n    # Fallback to the Course Objectives text block.\n    if not objectives:\n        m = re.search(\n            r"Course Objectives\\s*(.*?)(?:Program Learning Outcome|Course Learning Outcomes)",\n            text, re.I | re.S\n        )\n        if m:\n            for line in m.group(1).splitlines():\n                line = re.sub(r"^\\s*\\d+\\s*", "", clean(line))\n                line = re.sub(r"\\s+\\d+\\s*$", "", line)\n                if line and len(line.split()) >= 6:\n                    objectives.append(line)\n\n    # Exact CLO wording from tables.
    clos = {}
    for row in table_rows:
        if not row:
            continue
        cid = normalize_clo(row[0])
        if cid and len(row) >= 2:
            desc = clean(row[1])
            if desc and "course learning outcome" not in desc.lower():
                clos[cid] = desc

    # CLOs from paragraph/table text if necessary.
    for i, line in enumerate(all_lines):
        m = re.match(r"^(CLO\s*\d+)\s*(.*)$", line, re.I)
        if m:
            cid = normalize_clo(m.group(1))
            desc = clean(m.group(2))
            if not desc and i + 1 < len(all_lines):
                desc = clean(all_lines[i + 1])
            if desc:
                clos[cid] = desc

    def clo_sort_key(x):
        m = re.search(r"\d+", x)
        return int(m.group()) if m else 999

    clos = dict(sorted(clos.items(), key=lambda kv: clo_sort_key(kv[0])))
    return info, objectives, clos

def read_excel(data):
    return pd.read_excel(io.BytesIO(data), sheet_name=None, header=None)


# -----------------------------
# Excel parsers
# -----------------------------
def parse_original_obe_sheet(raw, sheet_name):
    """
    Parses the common OBE workbook layout used in the supplied SS1006 example:

    Row 1: course / teacher metadata
    Row 2: assessment labels + CLO Attainment group endings
    Row 3: dates + group maximum/category information
    Row 4: CLO.No
    Row 5: Weightage
    Row 6: Average
    Row 7: Std.Dev
    Row 8: Min
    Row 9: Max
    Row 10: student headers
    Row 11 onward: students

    The parser is intentionally based on structure, not hard-coded course values.
    """
    df = raw.copy()
    df = df.dropna(how="all")
    if df.empty:
        return None

    # Locate student-header row.
    header_idx = None
    for r in range(min(20, len(df))):
        vals = [clean(x) for x in df.iloc[r].tolist()]
        joined = " | ".join(vals).lower()
        if "rollno" in joined and "name" in joined and "sr." in joined:
            header_idx = r
            break

    # Locate the grouped OBE header row.
    group_header_idx = None
    for r in range(min(10, len(df))):
        vals = [clean(x) for x in df.iloc[r].tolist()]
        if sum("clo attainment" in x.lower() for x in vals) >= 1:
            group_header_idx = r
            break

    if header_idx is None or group_header_idx is None:
        return None

    group_header = [clean(x) for x in df.iloc[group_header_idx].tolist()]
    # Supporting rows are immediately below group header in the original structure.
    clo_no_row = min(group_header_idx + 2, len(df) - 1)
    weight_row = min(group_header_idx + 3, len(df) - 1)
    avg_row = min(group_header_idx + 4, len(df) - 1)
    sd_row = min(group_header_idx + 5, len(df) - 1)
    min_row = min(group_header_idx + 6, len(df) - 1)
    max_row = min(group_header_idx + 7, len(df) - 1)

    clo_no_vals = list(df.iloc[clo_no_row].tolist())
    weight_vals = list(df.iloc[weight_row].tolist())
    avg_vals = list(df.iloc[avg_row].tolist())
    sd_vals = list(df.iloc[sd_row].tolist())
    min_vals = list(df.iloc[min_row].tolist())
    max_vals = list(df.iloc[max_row].tolist())

    attainment_cols = [
        i for i, x in enumerate(group_header)
        if "clo attainment" in clean(x).lower()
    ]

    if not attainment_cols:
        return None

    # Build groups: all assessment columns since previous CLO Attainment column.
    groups = []
    prev_end = 4  # common student-information area; adjusted below if needed
    for end in attainment_cols:
        start = prev_end + 1
        if start >= end:
            start = max(0, end - 1)
        assessment_cols = list(range(start, end))
        # Remove columns that are clearly blank.
        assessment_cols = [c for c in assessment_cols if c < len(group_header) and clean(group_header[c])]
        groups.append((start, end, assessment_cols))
        prev_end = end

    # Better derive CLO id from CLO.No row inside each group.
    assessments = []
    clo_attainment_columns = {}
    for start, end, assessment_cols in groups:
        ids = []
        for c in assessment_cols:
            cid = normalize_clo(clo_no_vals[c])
            if cid:
                ids.append(cid)
            elif clean(clo_no_vals[c]).isdigit():
                ids.append(f"CLO{int(float(clean(clo_no_vals[c])))}")
        if ids:
            cid = ids[0]
        else:
            cid = f"CLO{len(clo_attainment_columns) + 1}"

        # In the original OBE layout, the column labelled "CLO Attainment"
        # is the raw CLO total; the following blank column stores the percentage.
        attainment_col = end + 1 if end + 1 < len(group_header) else end
        clo_attainment_columns[cid] = attainment_col

        for c in assessment_cols:
            label = clean(group_header[c])
            if not label:
                continue
            date_val = df.iloc[group_header_idx + 1, c] if group_header_idx + 1 < len(df) else ""
            assessments.append({
                "clo": cid,
                "assessment": label,
                "date": date_val,
                "weightage": num(weight_vals[c]) if c < len(weight_vals) else np.nan,
                "average": num(avg_vals[c]) if c < len(avg_vals) else np.nan,
                "sd": num(sd_vals[c]) if c < len(sd_vals) else np.nan,
                "minimum": num(min_vals[c]) if c < len(min_vals) else np.nan,
                "maximum": np.nan,
                "source": sheet_name
            })

    # Course metadata from the first row of the original OBE sheet.
    meta = {}
    first = [clean(x) for x in df.iloc[0].tolist()] if len(df) else []
    for i, val in enumerate(first):
        if val.lower() == "course" and i + 1 < len(first):
            meta["Course Title"] = clean(first[i + 1])
        if val.lower() == "teacher name" and i + 1 < len(first):
            meta["Instructor/Faculty"] = clean(first[i + 1])
    # The section is often stored as a standalone value in the first row.
    for val in first:
        if re.fullmatch(r"[A-Z]{2,10}-\d+[A-Z]\d+[A-Z]?", val):
            meta["Section"] = val
            break

    # Student table. The original OBE workbook uses numeric "Max" values
    # in the student-header row, so assessment/CLO column names are reconstructed
    # from the grouped header row instead of using those numeric cells.
    student_raw = df.iloc[header_idx + 1:].copy()

    base_headers = [clean(x) for x in df.iloc[header_idx].tolist()]
    headers = []
    for i in range(len(group_header)):
        if i < 5 and base_headers[i]:
            h = base_headers[i]
        elif i == len(group_header) - 1 and clean(group_header[i]).lower() == "g.tot":
            h = "Overall Score"
        else:
            h = clean(group_header[i])
            if not h:
                # Check whether this is a CLO attainment column.
                matched_cid = next((cid for cid, end_col in clo_attainment_columns.items() if end_col == i), None)
                h = f"{matched_cid} Attainment %" if matched_cid else f"Blank_{i+1}"
            elif i in attainment_cols:
                matched_cid = next((cid for cid, end_col in clo_attainment_columns.items() if end_col == i), None)
                h = f"{matched_cid} Attainment %" if matched_cid else h
            else:
                # Make repeated assessment labels unique enough for a dataframe.
                cid = next((cid for cid, end_col in clo_attainment_columns.items() if i < end_col and i > 0), "")
                h = f"{h} ({cid})" if cid else h

        if not h:
            h = f"Column_{i+1}"
        headers.append(h)

    # Ensure unique column names.
    seen = {}
    unique_headers = []
    for h in headers:
        seen[h] = seen.get(h, 0) + 1
        unique_headers.append(h if seen[h] == 1 else f"{h}_{seen[h]}")
    student_raw.columns = unique_headers

    sr_col = unique_headers[0]
    student_raw = student_raw[
        pd.to_numeric(student_raw[sr_col], errors="coerce").notna()
    ].copy()

    # Student CLO attainment columns now have explicit names.
    student_clo_cols = {}
    for cid, end_col in clo_attainment_columns.items():
        if end_col < len(unique_headers):
            student_clo_cols[cid] = unique_headers[end_col]

    # Build student/CLO data while retaining student identifiers for the output workbook.
    student_records = []
    for _, row in student_raw.iterrows():
        rec = {}
        for col in unique_headers[:5]:
            rec[clean(col)] = row[col]
        for cid, col in student_clo_cols.items():
            rec[f"{cid} Attainment %"] = num(row[col])
        if "Overall Score" in unique_headers:
            rec["Overall Score"] = num(row["Overall Score"])
        student_records.append(rec)

    student_df = pd.DataFrame(student_records)


    return {
        "meta": meta,
        "assessments": assessments,
        "student_df": student_df,
        "student_clo_cols": student_clo_cols,
        "format": "original_obe"
    }


def parse_normalized_sheet(raw, sheet_name):
    """Fallback parser for a conventional table with CLO1/CLO2/... columns."""
    df = raw.copy()
    df = df.dropna(how="all")
    if df.empty:
        return None

    # Try first 15 rows as headers.
    best = None
    for r in range(min(15, len(df))):
        vals = [clean(x) for x in df.iloc[r].tolist()]
        clo_cols = [x for x in vals if normalize_clo(x)]
        id_cols = [x for x in vals if re.search(r"student|roll|registration|reg\.?\s*no|name|id", x, re.I)]
        if clo_cols and id_cols:
            score = len(clo_cols) * 10 + len(id_cols) * 3
            if best is None or score > best[0]:
                best = (score, r, vals)

    if best is None:
        return None

    _, hr, headers0 = best
    headers = []
    seen = {}
    for i, h in enumerate(headers0):
        h = h or f"Column_{i+1}"
        seen[h] = seen.get(h, 0) + 1
        headers.append(h if seen[h] == 1 else f"{h}_{seen[h]}")

    t = df.iloc[hr + 1:].copy()
    t.columns = headers
    t = t.dropna(how="all").reset_index(drop=True)

    clo_cols = {normalize_clo(c): c for c in headers if normalize_clo(c)}
    id_cols = [c for c in headers if re.search(r"student|roll|registration|reg\.?\s*no|name|id", c, re.I)]

    if not clo_cols:
        return None

    assessments = []
    for cid, col in clo_cols.items():
        s = pd.to_numeric(t[col], errors="coerce")
        if s.notna().any():
            assessments.append({
                "clo": cid,
                "assessment": col,
                "date": "",
                "weightage": np.nan,
                "average": s.mean(),
                "sd": s.std(ddof=1),
                "minimum": s.min(),
                "maximum": s.max(),
                "source": sheet_name
            })

    student_df = pd.DataFrame()
    for c in id_cols:
        student_df[c] = t[c].tolist()
    for cid, col in clo_cols.items():
        student_df[f"{cid} Attainment %"] = pd.to_numeric(t[col], errors="coerce")
    overall_col = next(
        (c for c in headers if re.search(r"overall|total|aggregate|course", clean(c), re.I)),
        None
    )
    if overall_col:
        student_df["Overall Score"] = pd.to_numeric(t[overall_col], errors="coerce")
    else:
        att_cols = [c for c in student_df.columns if c.endswith("Attainment %")]
        if att_cols:
            student_df["Overall Score"] = student_df[att_cols].mean(axis=1)

    return {
        "meta": {},
        "assessments": assessments,
        "student_df": student_df,
        "student_clo_cols": {cid: f"{cid} Attainment %" for cid in clo_cols},
        "format": "normalized"
    }


def parse_obe_workbook(data):
    sheets = read_excel(data)

    # First priority: original/grouped OBE format.
    for name, raw in sheets.items():
        result = parse_original_obe_sheet(raw, name)
        if result is not None:
            return result

    # Fallback: conventional CLO-column table.
    for name, raw in sheets.items():
        result = parse_normalized_sheet(raw, name)
        if result is not None:
            return result

    raise ValueError(
        "The workbook format could not be recognized. Upload the original OBE assessment Excel workbook or a student table containing CLO-labelled attainment columns."
    )


# -----------------------------
# Analysis
# -----------------------------
def analyze(obe, clos):
    assessments = obe["assessments"]
    student_df = obe["student_df"].copy()

    # If Excel has CLOs not present in the outline, flag them.
    excel_clos = []
    for a in assessments:
        if a["clo"] not in excel_clos:
            excel_clos.append(a["clo"])
    for c in obe["student_clo_cols"]:
        if c not in excel_clos:
            excel_clos.append(c)

    # Add unmatched Excel CLOs to analysis so discrepancy is visible.
    all_clos = list(clos.keys())
    for c in excel_clos:
        if c not in all_clos:
            clos[c] = "Excel assessment CLO has no matching Course Outline CLO."
            all_clos.append(c)

    stats = {}
    for cid in all_clos:
        col = obe["student_clo_cols"].get(cid)
        if col is not None and col in student_df.columns:
            s = pd.to_numeric(student_df[col], errors="coerce").dropna()
        else:
            s = pd.Series(dtype=float)

        if len(s):
            stats[cid] = {
                "n": int(len(s)),
                "mean": float(s.mean()),
                "sd": float(s.std(ddof=1)) if len(s) > 1 else np.nan,
                "n_benchmark": int((s >= BENCHMARK).sum()),
                "pct_benchmark": float((s >= BENCHMARK).mean() * 100),
                "evidence": [a["assessment"] for a in assessments if a["clo"] == cid],
            }
        else:
            stats[cid] = {
                "n": 0,
                "mean": np.nan,
                "sd": np.nan,
                "n_benchmark": 0,
                "pct_benchmark": np.nan,
                "evidence": [a["assessment"] for a in assessments if a["clo"] == cid],
            }

    # Overall score.
    if "Overall Score" in student_df.columns:
        overall = pd.to_numeric(student_df["Overall Score"], errors="coerce").dropna()
    else:
        att_cols = [c for c in student_df.columns if c.endswith("Attainment %")]
        overall = student_df[att_cols].mean(axis=1, skipna=True).dropna() if att_cols else pd.Series(dtype=float)

    overall_stats = {
        "n": int(len(overall)),
        "highest": float(overall.max()) if len(overall) else np.nan,
        "lowest": float(overall.min()) if len(overall) else np.nan,
        "mean": float(overall.mean()) if len(overall) else np.nan,
        "median": float(overall.median()) if len(overall) else np.nan,
        "sd": float(overall.std(ddof=1)) if len(overall) > 1 else np.nan,
        "benchmark_pct": float((overall >= BENCHMARK).mean() * 100) if len(overall) else np.nan
    }

    return clos, stats, overall_stats, student_df, assessments


# -----------------------------
# Charts
# -----------------------------
def make_charts(stats, assessments, output_dir):
    output_dir = Path(output_dir)
    output_dir.mkdir(exist_ok=True)

    paths = []
    clo_ids = list(stats.keys())

    # 1. CLO attainment
    vals = [stats[c]["mean"] for c in clo_ids]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.bar(clo_ids, [0 if pd.isna(v) else v for v in vals])
    ax.axhline(BENCHMARK, linestyle="--", linewidth=1.5, label=f"{BENCHMARK:.0f}% benchmark")
    ax.set_title("Figure 1. CLO-wise OBE Attainment")
    ax.set_xlabel("Course Learning Outcome")
    ax.set_ylabel("Mean Attainment (%)")
    ax.set_ylim(0, 100)
    for i, v in enumerate(vals):
        if not pd.isna(v):
            ax.text(i, min(v + 2, 97), f"{v:.2f}%", ha="center")
    ax.legend()
    fig.tight_layout()
    p = output_dir / "CLO_Attainment_Chart.png"
    fig.savefig(p, dpi=200, bbox_inches="tight")
    plt.close(fig)
    paths.append(p)

    # 2. Benchmark achievement
    vals = [stats[c]["pct_benchmark"] for c in clo_ids]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.bar(clo_ids, [0 if pd.isna(v) else v for v in vals])
    ax.set_title(f"Figure 2. Students Achieving ≥{BENCHMARK:.0f}% by CLO")
    ax.set_xlabel("Course Learning Outcome")
    ax.set_ylabel("Students achieving benchmark (%)")
    ax.set_ylim(0, 100)
    for i, v in enumerate(vals):
        if not pd.isna(v):
            ax.text(i, min(v + 2, 97), f"{v:.0f}%", ha="center")
    fig.tight_layout()
    p = output_dir / "Benchmark_Achievement_Chart.png"
    fig.savefig(p, dpi=200, bbox_inches="tight")
    plt.close(fig)
    paths.append(p)

    # 3. Assessment performance
    valid = [a for a in assessments if not pd.isna(a["average"])]
    vals = [a["average"] for a in valid]
    labels = [f"{a['clo']}\n{a['assessment']}" for a in valid]
    fig, ax = plt.subplots(figsize=(12, 6))
    ax.bar(range(len(vals)), vals)
    ax.set_title("Figure 3. Assessment Mean Scores as Reported in Excel")
    ax.set_xlabel("Assessment / CLO")
    ax.set_ylabel("Mean score")
    ax.set_xticks(range(len(vals)))
    ax.set_xticklabels(labels, rotation=55, ha="right", fontsize=8)
    fig.tight_layout()
    p = output_dir / "Assessment_Performance_Chart.png"
    fig.savefig(p, dpi=200, bbox_inches="tight")
    plt.close(fig)
    paths.append(p)

    return paths


# -----------------------------
# Excel output
# -----------------------------
def make_excel(info, clos, assessments, stats, overall, student_df, output):
    wb = Workbook()
    wb.remove(wb.active)

    title_fill = PatternFill("solid", fgColor="1F4E78")
    header_fill = PatternFill("solid", fgColor="D9EAF7")

    def setup(ws, title, headers):
        ws.merge_cells(
            start_row=1, start_column=1,
            end_row=1, end_column=max(1, len(headers))
        )
        cell = ws.cell(1, 1, title)
        cell.fill = title_fill
        cell.font = Font(color="FFFFFF", bold=True, size=12)
        for j, h in enumerate(headers, 1):
            c = ws.cell(3, j, h)
            c.fill = header_fill
            c.font = Font(bold=True)
            c.alignment = Alignment(wrap_text=True, vertical="top")
        ws.freeze_panes = "A4"

    def put(ws, rows):
        for r, row in enumerate(rows, 4):
            for j, value in enumerate(row, 1):
                if isinstance(value, np.generic):
                    value = value.item()
                ws.cell(r, j, value)

    ws = wb.create_sheet("OBE Summary")
    headers = ["CLO", "Official CLO Description", "Mean Attainment (%)",
               f"Students ≥{BENCHMARK:.0f}%", f"Students ≥{BENCHMARK:.0f}% (%)",
               "SD", "Status"]
    setup(ws, "OBE Summary", headers)
    put(ws, [
        [
            c, clos[c], stats[c]["mean"], stats[c]["n_benchmark"],
            stats[c]["pct_benchmark"], stats[c]["sd"], status(stats[c]["mean"])
        ]
        for c in clos
    ])

    ws = wb.create_sheet("CLO–Assessment Mapping")
    headers = ["CLO", "Official CLO Description", "Assessment/Question",
               "Date", "Weightage", "Average/Mean Score",
               "Maximum Marks", "Attainment %", "Source"]
    setup(ws, "CLO–Assessment Mapping", headers)
    rows = []
    for a in assessments:
        # Raw maximum marks are reported only when present in the input.
        maximum = a.get("maximum", np.nan)
        maximum_out = maximum if not pd.isna(maximum) else MISSING
        rows.append([
            a["clo"], clos.get(a["clo"], MISSING), a["assessment"],
            a.get("date", ""), a.get("weightage", np.nan),
            a.get("average", np.nan), maximum_out, MISSING, a["source"]
        ])
    put(ws, rows)

    ws = wb.create_sheet("Assessment Analysis")
    headers = ["CLO", "Assessment/Question", "Weightage",
               "Average/Mean Score", "Assessment Attainment %",
               "Interpretation"]
    setup(ws, "Assessment Analysis", headers)
    rows = []
    for a in assessments:
        rows.append([
            a["clo"], a["assessment"], a.get("weightage", np.nan),
            a.get("average", np.nan), MISSING,
            f"Mean score = {a.get('average', np.nan):.2f}; assessment-level percentage is not calculated unless a raw maximum is explicitly available."
            if not pd.isna(a.get("average", np.nan))
            else MISSING
        ])
    put(ws, rows)

    ws = wb.create_sheet("Student/CLO Data")
    student_headers = list(student_df.columns)
    setup(ws, "Student/CLO Data", student_headers)
    put(ws, student_df.values.tolist())

    ws = wb.create_sheet("CQI Action Plan")
    headers = ["CLO/Area", "Identified Issue", "Recommended Action",
               "Teaching/Learning Intervention", "Follow-up Evidence", "Target"]
    setup(ws, "CQI Action Plan", headers)
    rows = []
    for c in clos:
        mean = stats[c]["mean"]
        if pd.isna(mean):
            issue = MISSING
            action = f"Review evidence for the exact official CLO: {clos[c]}"
        elif mean < BENCHMARK:
            issue = f"Mean attainment {mean:.2f}% is below the {BENCHMARK:.0f}% benchmark; {stats[c]['pct_benchmark']:.0f}% achieved the benchmark."
            action = f"Strengthen learning and assessment activities aligned to: {clos[c]}"
        else:
            issue = f"Mean attainment {mean:.2f}% meets/exceeds the {BENCHMARK:.0f}% benchmark."
            action = f"Maintain and monitor practices aligned to: {clos[c]}"
        rows.append([
            c, issue, action,
            "Use CLO-aligned practice, formative assessment, guided application and feedback.",
            f"Repeat CLO-aligned checks and compare mean attainment and ≥{BENCHMARK:.0f}% achievement.",
            f"Mean attainment ≥{BENCHMARK:.0f}%"
        ])
    put(ws, rows)

    ws = wb.create_sheet("Chart Data")
    headers = ["CLO", "Mean Attainment (%)",
               f"Students ≥{BENCHMARK:.0f}% (%)", "Status"]
    setup(ws, "Chart Data", headers)
    put(ws, [
        [c, stats[c]["mean"], stats[c]["pct_benchmark"], status(stats[c]["mean"])]
        for c in clos
    ])

    # Reasonable widths.
    for ws in wb.worksheets:
        for col in ws.columns:
            letter = col[0].column_letter
            ws.column_dimensions[letter].width = min(
                max(12, max(len(clean(c.value)) for c in col) + 2), 55
            )

    wb.save(output)


# -----------------------------
# Word report
# -----------------------------
def make_docx(info, objectives, clos, assessments, stats, overall, student_df, chart_paths, output):
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.70)
    sec.right_margin = Inches(0.70)

    # Default font.
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10)

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = str(h)
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        return t

    # Cover page.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OUTCOME-BASED EDUCATION (OBE)\n")
    r.bold = True
    r.font.size = Pt(18)
    r = p.add_run("EVALUATION REPORT\n\n")
    r.bold = True
    r.font.size = Pt(22)
    p.add_run(f"{safe(info.get('Course Title'))} ({safe(info.get('Course Code'))})\n")
    p.add_run(f"Program: {safe(info.get('Program'))}\n")
    p.add_run(f"Section: {safe(info.get('Section'))}\n")
    p.add_run(f"Semester: {safe(info.get('Semester'))}\n")
    p.add_run(f"Prepared from uploaded Course Outline and OBE assessment workbook")
    doc.add_page_break()

    valid = {c: stats[c]["mean"] for c in clos if not pd.isna(stats[c]["mean"])}
    strongest = max(valid, key=valid.get) if valid else None
    weakest = min(valid, key=valid.get) if valid else None
    overall_mean = overall["mean"]

    # 1 Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    exec_rows = [
        ["Course", f"{safe(info.get('Course Title'))} ({safe(info.get('Course Code'))})"],
        ["Program / Section", f"{safe(info.get('Program'))} / {safe(info.get('Section'))}"],
        ["Semester", safe(info.get("Semester"))],
        ["Number of students", overall["n"]],
        ["Overall CLO/course attainment", f"{overall_mean:.2f}%" if not pd.isna(overall_mean) else MISSING],
        ["Strongest CLO", f"{strongest} – {valid[strongest]:.2f}%" if strongest else MISSING],
        ["Weakest CLO", f"{weakest} – {valid[weakest]:.2f}%" if weakest else MISSING],
        [f"CLOs ≥{BENCHMARK:.0f}%", sum(1 for c in clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] >= BENCHMARK)],
        [f"CLOs <{BENCHMARK:.0f}%", sum(1 for c in clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] < BENCHMARK)]
    ]
    table(["Item", "Result"], exec_rows)

    # 2 Course information
    doc.add_heading("2. Course Information", level=1)
    table(["Field", "Information"], [[k, safe(v)] for k, v in info.items() if k != "Course Description"])

    doc.add_heading("2.1 Course Description", level=2)
    doc.add_paragraph(safe(info.get("Course Description")))

    doc.add_heading("2.2 Course Objectives", level=2)
    if objectives:
        for i, obj in enumerate(objectives, 1):
            doc.add_paragraph(f"{i}. {obj}")
    else:
        doc.add_paragraph(MISSING)

    doc.add_heading("2.3 Official CLOs", level=2)
    table(["CLO", "Official CLO Statement"], [[c, clos[c]] for c in clos])

    # 3 Alignment
    doc.add_heading("3. CLO–Assessment Alignment", level=1)
    rows = []
    for a in assessments:
        maximum = a.get("maximum", np.nan)
        maximum_out = maximum if not pd.isna(maximum) else MISSING
        rows.append([
            a["clo"], clos.get(a["clo"], MISSING), a["assessment"],
            maximum_out, MISSING
        ])
    table(["CLO", "Official CLO", "Assessment/Question", "Maximum Marks", "Attainment %"], rows)

    # 4 Methodology
    doc.add_heading("4. Methodology", level=1)
    doc.add_paragraph(
        "The Course Outline is used as the authoritative source for course information and exact official CLO wording. "
        "The uploaded OBE Excel workbook is used as the authoritative source for student performance, assessment evidence, "
        "marks and CLO attainment. Missing information is reported rather than inferred. "
        f"The benchmark used in this report is {BENCHMARK:.0f}%."
    )

    # 5 CLO attainment
    doc.add_heading("5. CLO-wise OBE Attainment", level=1)
    rows = []
    for c in clos:
        s = stats[c]
        rows.append([
            c, clos[c],
            f"{s['mean']:.2f}%" if not pd.isna(s["mean"]) else MISSING,
            f"{s['n_benchmark']} ({s['pct_benchmark']:.0f}%)" if not pd.isna(s["pct_benchmark"]) else MISSING,
            status(s["mean"])
        ])
    table(["CLO", "CLO Description", "Mean Attainment (%)",
           f"Students ≥{BENCHMARK:.0f}%", "Status"], rows)

    if valid:
        doc.add_paragraph(
            f"Overall weighted course/CLO attainment based on the uploaded student-level overall score is "
            f"{overall_mean:.2f}%." if not pd.isna(overall_mean) else ""
        )

    # 6 Assessment analysis
    doc.add_heading("6. Assessment-wise Analysis", level=1)
    rows = []
    for a in assessments:
        avg = a.get("average", np.nan)
        rows.append([
            a["clo"], a["assessment"],
            a.get("date", ""),
            a.get("weightage", ""),
            f"{avg:.2f}" if not pd.isna(avg) else MISSING,
            MISSING
        ])
    table(["CLO", "Assessment/Question", "Date", "Weightage",
           "Average/Mean Score", "Attainment %"], rows)
    doc.add_paragraph(
        "Assessment means are reported on the raw scales supplied by the Excel workbook. "
        "Assessment-level attainment percentages are not calculated when a separate raw maximum mark is not explicitly available."
    )

    # 7 Student performance
    doc.add_heading("7. Student Performance Analysis", level=1)
    rows = [
        ["Number of students assessed", overall["n"]],
        ["Highest overall score", f"{overall['highest']:.2f}" if not pd.isna(overall["highest"]) else MISSING],
        ["Lowest overall score", f"{overall['lowest']:.2f}" if not pd.isna(overall["lowest"]) else MISSING],
        ["Mean overall score", f"{overall['mean']:.2f}" if not pd.isna(overall["mean"]) else MISSING],
        ["Median", f"{overall['median']:.2f}" if not pd.isna(overall["median"]) else MISSING],
        ["Standard deviation", f"{overall['sd']:.2f}" if not pd.isna(overall["sd"]) else MISSING],
        [f"Percentage meeting {BENCHMARK:.0f}% benchmark",
         f"{overall['benchmark_pct']:.0f}%" if not pd.isna(overall["benchmark_pct"]) else MISSING]
    ]
    table(["Metric", "Result"], rows)

    # Distribution
    if "Overall Score" in student_df.columns:
        scores = pd.to_numeric(student_df["Overall Score"], errors="coerce").dropna()
        bands = [
            ("<60%", (scores < 60).sum()),
            ("60–69.99%", ((scores >= 60) & (scores < 70)).sum()),
            ("70–79.99%", ((scores >= 70) & (scores < 80)).sum()),
            ("80–89.99%", ((scores >= 80) & (scores < 90)).sum()),
            ("90–100%", (scores >= 90).sum())
        ]
        table(
            ["Overall-score band", "Students", "Percentage"],
            [[b, int(n), f"{(n/len(scores))*100:.1f}%"] for b, n in bands] if len(scores) else []
        )

    # 8 Charts
    doc.add_heading("8. Charts and Visual Evidence", level=1)
    for i, pth in enumerate(chart_paths, 1):
        doc.add_picture(str(pth), width=Inches(6.5))
        q = doc.add_paragraph(
            f"Figure {i}. {pth.stem.replace('_', ' ')}."
        )
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 9 CLO alignment
    doc.add_heading("9. CLO Alignment with Course Outline", level=1)
    rows = []
    for c in clos:
        s = stats[c]
        evidence = "; ".join(s["evidence"]) if s["evidence"] else NO_EVIDENCE
        priority = "High" if not pd.isna(s["mean"]) and s["mean"] < BENCHMARK else "Maintain"
        rows.append([
            c, clos[c], evidence,
            f"{s['mean']:.2f}%" if not pd.isna(s["mean"]) else MISSING,
            f"{s['pct_benchmark']:.0f}%" if not pd.isna(s["pct_benchmark"]) else MISSING,
            status(s["mean"]), priority
        ])
    table(
        ["CLO", "Official CLO", "Assessment Evidence", "Attainment %",
         f"Benchmark Achievement %", "Status", "CQI Priority"], rows
    )

    # 10 Findings
    doc.add_heading("10. Key Findings and OBE Interpretation", level=1)
    if valid:
        above = [c for c in clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] >= BENCHMARK]
        below = [c for c in clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] < BENCHMARK]
        doc.add_paragraph(
            f"Overall student-level course score mean is {overall_mean:.2f}%."
            if not pd.isna(overall_mean) else
            "An overall course mean was not available."
        )
        doc.add_paragraph(
            f"{len(above)} of {len(valid)} CLOs meet or exceed the {BENCHMARK:.0f}% benchmark."
        )
        if strongest:
            doc.add_paragraph(f"{strongest} has the highest observed CLO mean attainment ({valid[strongest]:.2f}%).")
        if weakest:
            doc.add_paragraph(f"{weakest} has the lowest observed CLO mean attainment ({valid[weakest]:.2f}%).")
        if below:
            doc.add_paragraph(
                "The CLO(s) below benchmark provide evidence for targeted CQI attention: " +
                ", ".join(below) + "."
            )
    doc.add_paragraph(
        "These findings are evidence-based descriptions of the supplied assessment data and do not by themselves establish claims about instructor effectiveness, teaching quality or student ability."
    )

    # 11 CQI
    doc.add_heading("11. CQI / Action Plan", level=1)
    rows = []
    for c in clos:
        s = stats[c]
        mean = s["mean"]
        if pd.isna(mean):
            issue = MISSING
            action = f"Review assessment evidence for {c}: {clos[c]}"
        elif mean < BENCHMARK:
            issue = f"Mean attainment {mean:.2f}% is below {BENCHMARK:.0f}%; {s['pct_benchmark']:.0f}% achieved the benchmark."
            action = f"Strengthen explicit instruction and practice aligned to the official CLO: {clos[c]}"
        else:
            issue = f"Mean attainment {mean:.2f}% meets/exceeds {BENCHMARK:.0f}%."
            action = f"Maintain and document effective practices aligned to: {clos[c]}"
        rows.append([
            c, issue, action,
            "Use CLO-aligned concept/application tasks, guided practice, formative assessment and feedback.",
            f"Repeat CLO-aligned formative checks and compare mean attainment and ≥{BENCHMARK:.0f}% achievement.",
            f"Mean attainment ≥{BENCHMARK:.0f}%"
        ])
    table(
        ["CLO/Area", "Identified Issue", "Recommended Action",
         "Teaching/Learning Intervention", "Follow-up Evidence", "Target"], rows
    )

    # 12 Conclusion
    doc.add_heading("12. Formal Conclusion", level=1)
    if not pd.isna(overall_mean):
        doc.add_paragraph(
            f"The OBE evidence for {safe(info.get('Course Title'))} ({safe(info.get('Course Code'))}) "
            f"indicates an overall course/student-score mean of {overall_mean:.2f}%. "
            f"The CLO-level analysis identifies the observed strongest and weakest outcomes and provides "
            f"a targeted CQI action plan for outcomes below the {BENCHMARK:.0f}% benchmark."
        )
    else:
        doc.add_paragraph(
            "The uploaded files did not provide a usable overall course-score measure. "
            "CLO-level evidence is reported where available."
        )

    # 13 Quality check
    doc.add_heading("13. OBE Quality Check", level=1)
    outline_clos = [c for c in clos if not clos[c].startswith("Excel assessment CLO")]
    matched = [c for c in outline_clos if c in excel_clos]
    unmatched_excel = [c for c in excel_clos if c not in outline_clos]
    rows = [
        ["Number of students analyzed", overall["n"]],
        ["Overall CLO/course attainment", f"{overall_mean:.2f}%" if not pd.isna(overall_mean) else MISSING],
        ["Strongest CLO", f"{strongest} ({valid[strongest]:.2f}%)" if strongest else MISSING],
        ["Weakest CLO", f"{weakest} ({valid[weakest]:.2f}%)" if weakest else MISSING],
        [f"CLOs ≥{BENCHMARK:.0f}%", sum(1 for c in outline_clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] >= BENCHMARK)],
        [f"CLOs <{BENCHMARK:.0f}%", sum(1 for c in outline_clos if not pd.isna(stats[c]["mean"]) and stats[c]["mean"] < BENCHMARK)],
        ["Assessments analyzed", len(assessments)],
        ["CLOs matched between files", f"{len(matched)} of {len(outline_clos)}"],
        ["Excel CLOs without Course Outline match", ", ".join(unmatched_excel) if unmatched_excel else "None identified"]
    ]
    table(["Quality Check", "Result"], rows)

    doc.add_paragraph(
        "Data limitation: where the Excel workbook did not separately specify raw maximum marks for an assessment component, "
        "maximum marks and assessment-level attainment percentages were not inferred."
    )

    doc.save(output)


# -----------------------------
# UI
# -----------------------------
st.markdown('<div class="section-title">1. Course Information</div>', unsafe_allow_html=True)
st.caption("Enter details manually, or upload the Course Outline below and the fields will be populated from it.")

c1, c2, c3 = st.columns(3)
with c1:
    institution = st.text_input("Institution", value="")
    department = st.text_input("Department", value="")
    program = st.text_input("Program", value="")
with c2:
    course_title = st.text_input("Course Title", value="")
    course_code = st.text_input("Course Code", value="")
    semester = st.text_input("Semester", value="")
with c3:
    academic_year = st.text_input("Academic Year", value="")
    instructor = st.text_input("Course Teacher / Instructor", value="")
    credit_hours = st.text_input("Credit Hours", value="")

section = st.text_input("Section (optional)", value="")
campus = st.text_input("Campus (optional)", value="")

st.markdown('<div class="section-title">2. Course Learning Outcomes</div>', unsafe_allow_html=True)
outline = st.file_uploader(
    "Upload Course Outline (.docx)",
    type=["docx"],
    key="outline_upload"
)

if outline:
    try:
        parsed_info, objectives, parsed_clos = parse_outline(outline.getvalue())
        # Populate blank fields without overwriting manual entries.
        mapping = {
            "Institution": ("institution", parsed_info.get("Institution")),
            "Department": ("department", parsed_info.get("Department")),
            "Program": ("program", parsed_info.get("Program")),
            "Course Title": ("course_title", parsed_info.get("Course Title")),
            "Course Code": ("course_code", parsed_info.get("Course Code")),
            "Semester": ("semester", parsed_info.get("Semester")),
            "Academic Year": ("academic_year", parsed_info.get("Academic Year")),
            "Instructor/Faculty": ("instructor", parsed_info.get("Instructor/Faculty")),
            "Credit Hours": ("credit_hours", parsed_info.get("Credit Hours")),
            "Section": ("section", parsed_info.get("Section")),
            "Campus": ("campus", parsed_info.get("Campus"))
        }
        st.success("Course Outline loaded. Official CLOs are extracted from the uploaded outline.")
        if parsed_clos:
            st.dataframe(
                pd.DataFrame(
                    [{"CLO": c, "Official CLO Description": d} for c, d in parsed_clos.items()]
                ),
                use_container_width=True,
                hide_index=True
            )
        else:
            st.warning("No CLOs were detected automatically. Enter the official CLO wording manually below.")
    except Exception as e:
        parsed_info, objectives, parsed_clos = {}, [], {}
        st.error(f"Could not read the Course Outline: {e}")
else:
    parsed_info, objectives, parsed_clos = {}, [], {}

# Manual/override fields for missing CLOs.
clos = dict(parsed_clos)
if not clos:
    st.info("You can enter official CLO wording manually if a Course Outline is unavailable or unreadable.")
    n_clo = st.number_input("Number of CLOs", 1, 20, 5, 1)
    for i in range(1, n_clo + 1):
        clos[f"CLO{i}"] = st.text_area(
            f"CLO {i} official wording",
            value="",
            key=f"manual_clo_{i}"
        )
    clos = {k: v for k, v in clos.items() if clean(v)}

st.markdown('<div class="section-title">3. Student Assessment Data</div>', unsafe_allow_html=True)
excel = st.file_uploader(
    "Upload OBE Assessment Excel (.xlsx)",
    type=["xlsx"],
    key="excel_upload"
)

if excel:
    try:
        obe = parse_obe_workbook(excel.getvalue())
        st.success(
            f"Excel loaded successfully. Detected {len(obe['student_df'])} student records and "
            f"{len(obe['assessments'])} assessment components from the '{obe['format']}' structure."
        )
    except Exception as e:
        obe = None
        st.error(f"Could not analyze the uploaded Excel workbook: {e}")
else:
    obe = None

# Build final info from manual fields first, then outline values.
info = {
    "Institution": institution or parsed_info.get("Institution", ""),
    "Department": department or parsed_info.get("Department", ""),
    "Program": program or parsed_info.get("Program", ""),
    "Course Title": course_title or parsed_info.get("Course Title", ""),
    "Course Code": course_code or parsed_info.get("Course Code", ""),
    "Semester": semester or parsed_info.get("Semester", ""),
    "Academic Year": academic_year or parsed_info.get("Academic Year", ""),
    "Campus": campus or parsed_info.get("Campus", ""),
    "Instructor/Faculty": instructor or parsed_info.get("Instructor/Faculty", ""),
    "Credit Hours": credit_hours or parsed_info.get("Credit Hours", ""),
    "Section": section or parsed_info.get("Section", ""),
    "Course Description": parsed_info.get("Course Description", "")
}

# Add workbook metadata where outline/manual fields are empty.
if obe:
    for k, v in obe.get("meta", {}).items():
        if not info.get(k):
            info[k] = v

if obe and clos:
    try:
        clos, stats, overall, student_df, assessments = analyze(obe, clos)

        st.markdown('<div class="section-title">4. OBE Analysis Dashboard</div>', unsafe_allow_html=True)

        cards = [
            ("Institution", safe(info.get("Institution"))),
            ("Department", safe(info.get("Department"))),
            ("Program", safe(info.get("Program"))),
            ("Course Title", safe(info.get("Course Title"))),
            ("Course Code", safe(info.get("Course Code"))),
            ("Semester", safe(info.get("Semester"))),
            ("Academic Year", safe(info.get("Academic Year"))),
            ("Course Teacher / Instructor", safe(info.get("Instructor/Faculty"))),
            ("Credit Hours", safe(info.get("Credit Hours")))
        ]

        for start in range(0, len(cards), 3):
            cols = st.columns(3)
            for col, (label, value) in zip(cols, cards[start:start + 3]):
                with col:
                    st.markdown(
                        f'<div class="info-card"><div class="info-label">{label}</div><div class="info-value">{value}</div></div>',
                        unsafe_allow_html=True
                    )

        st.markdown("### CLO Attainment")
        result_df = pd.DataFrame([
            {
                "CLO": c,
                "Exact Official CLO": clos[c],
                "Mean Attainment (%)": stats[c]["mean"],
                f"Students ≥{BENCHMARK:.0f}%": stats[c]["n_benchmark"],
                f"Students ≥{BENCHMARK:.0f}% (%)": stats[c]["pct_benchmark"],
                "SD": stats[c]["sd"],
                "Status": status(stats[c]["mean"])
            }
            for c in clos
        ])
        st.dataframe(result_df, use_container_width=True, hide_index=True)

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Students Assessed", overall["n"])
        m2.metric("Overall %", f"{overall['mean']:.2f}%" if not pd.isna(overall["mean"]) else "N/A")

        valid = {c: stats[c]["mean"] for c in clos if not pd.isna(stats[c]["mean"])}
        if valid:
            strongest = max(valid, key=valid.get)
            weakest = min(valid, key=valid.get)
            m3.metric("Strongest CLO", f"{strongest} ({valid[strongest]:.2f}%)")
            m4.metric("Weakest CLO", f"{weakest} ({valid[weakest]:.2f}%)")

        st.markdown("### Assessment Evidence")
        assessment_view = pd.DataFrame([
            {
                "CLO": a["clo"],
                "Assessment/Question": a["assessment"],
                "Date": a.get("date", ""),
                "Weightage": a.get("weightage", np.nan),
                "Average/Mean Score": a.get("average", np.nan),
                "Maximum Marks": a.get("maximum", np.nan) if not pd.isna(a.get("maximum", np.nan)) else MISSING
            }
            for a in assessments
        ])
        st.dataframe(assessment_view, use_container_width=True, hide_index=True)

        st.markdown("### Student Performance")
        if "Overall Score" in student_df.columns:
            st.dataframe(
                student_df.head(100),
                use_container_width=True,
                hide_index=True
            )

        # Generate package.
        if st.button("📦 Generate Complete OBE Report Package", type="primary", use_container_width=True):
            out = Path("obe_output")
            out.mkdir(exist_ok=True)

            chart_paths = make_charts(stats, assessments, out)
            xlsx_path = out / "OBE_Analysis.xlsx"
            docx_path = out / "OBE_Evaluation_Report.docx"

            make_excel(
                info, clos, assessments, stats, overall, student_df, xlsx_path
            )
            make_docx(
                info, objectives, clos, assessments, stats, overall,
                student_df, chart_paths, docx_path
            )

            st.success("Complete OBE package generated successfully.")

            st.markdown("### Download Files")
            st.download_button(
                "📄 Download Word OBE Evaluation Report",
                docx_path.read_bytes(),
                file_name=docx_path.name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            st.download_button(
                "📊 Download Excel OBE Analysis Workbook",
                xlsx_path.read_bytes(),
                file_name=xlsx_path.name,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )

            for p in chart_paths:
                st.download_button(
                    f"📈 Download {p.name}",
                    p.read_bytes(),
                    file_name=p.name,
                    mime="image/png",
                    use_container_width=True
                )

            st.markdown("### Chart Preview")
            for p in chart_paths:
                st.image(str(p), caption=p.stem, use_container_width=True)

    except Exception as e:
        st.error(f"OBE analysis could not be completed: {e}")

else:
    if not clos:
        st.info("Enter or upload the official CLOs to continue.")
    elif not excel:
        st.info("Upload the OBE assessment Excel workbook to generate the dynamic OBE analysis.")
    elif obe is None:
        st.info("Please upload a readable OBE Excel workbook.")

# -----------------------------
# Master Prompt
# -----------------------------
MASTER_PROMPT = """OBE EVALUATION REPORT – MASTER PROMPT

Analyze the uploaded Course Outline and OBE Excel workbook together.

SOURCE HIERARCHY
Course Outline = authoritative for course information and exact official CLO wording.
Excel = authoritative for numerical calculations, student performance, marks, assessment results, assessment-to-CLO mapping and attainment.
Never invent, estimate, modify or reinterpret values. Never change official CLO wording. Never infer missing maximum marks. Missing information must be reported as “Not available in the provided files.”

REQUIRED OUTPUT
1. Course Information: code, title, credit hours, program, semester, campus, instructor, description, objectives, official CLOs, CLO/PLO mapping where available.
2. CLO–Assessment Alignment: CLO, exact official CLO, assessment/question, maximum marks, source.
3. OBE Analysis: CLO mean attainment, SD, number and percentage ≥70%, overall attainment, assessment performance, strongest/weakest CLO, CLOs above/below benchmark.
4. Status: ≥80 Strong; 70–79.99 Satisfactory; <70 Needs Improvement.
5. CLO Attainment Table: CLO | exact CLO description | mean attainment | students ≥70% | status.
6. Assessment-to-CLO Mapping: CLO | assessment/question | maximum marks | average/mean | attainment % | interpretation. If maximum marks are absent, do not infer attainment %.
7. Student Performance: N, highest, lowest, mean, median, SD, distribution, benchmark percentage. Avoid unnecessary personal data.
8. Charts: CLO attainment with 70% line; benchmark achievement; assessment performance. Include titles, axes, percentages where appropriate and figure captions.
9. Evidence-based OBE interpretation. Do not make unsupported claims about student ability, instructor effectiveness or teaching quality.
10. CQI Action Plan: specific to each weak CLO and its exact official wording; include issue, action, intervention, follow-up evidence and target.
11. Dedicated “CLO Alignment with Course Outline” section: CLO | Official CLO | Assessment Evidence | Attainment % | Benchmark Achievement % | Status | CQI Priority.
12. Formal conclusion for HOD/OBE Committee/QEC/accreditation records.
13. Executive Summary at the beginning.
14. Word report with cover page, executive summary, course information, description, objectives, official CLOs, alignment, methodology, attainment, assessment analysis, student analysis, charts, CLO alignment, findings, CQI, conclusion and quality check.
15. Separate Excel workbook with sheets: OBE Summary; CLO–Assessment Mapping; Assessment Analysis; Student/CLO Data; CQI Action Plan; Chart Data.
16. Output DOCX, XLSX and three PNG charts.
17. Validate student count, mappings, marks, means, attainment, benchmark percentages, exact CLO wording, and consistency across tables/charts.
18. Final quality check: students, overall attainment, strongest/weakest CLO, CLOs ≥70 and <70, assessments, matched CLOs, discrepancies, internal consistency.

If a Course Outline CLO has no Excel evidence, write exactly: “No assessment evidence for this CLO was identified in the provided Excel file.”
If an Excel assessment CLO cannot be matched to the Course Outline, flag the discrepancy clearly.
"""

with st.expander("View / download the OBE Master Prompt"):
    st.text_area("Master Prompt", MASTER_PROMPT, height=420)
    st.download_button(
        "Download OBE Master Prompt",
        MASTER_PROMPT,
        file_name="OBE_Evaluation_Master_Prompt.txt",
        mime="text/plain"
    )
